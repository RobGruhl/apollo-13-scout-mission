#!/usr/bin/env python3
"""
Generate a comprehensive Apollo 13 Word document from content files and images.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path
import re

def create_apollo13_document():
    """Create the complete Apollo 13 Word document."""
    doc = Document()

    # Set up document
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add title page
    add_title_page(doc)

    # Content mapping: (content_file, image_files)
    # Based on presentation order from README
    content_order = [
        # INTRO
        ('narrative/01-launch-and-mission-overview.txt', ['content/images/slide01_launch.jpg']),
        ('info/01-meet-the-crew.txt', ['content/images/slide02_meet_the_crew.jpg']),

        # CRISIS BEGINS
        ('narrative/02-the-explosion.txt', ['content/images/slide03_explosion.jpg']),
        ('info/02-what-caused-the-explosion.txt', ['content/images/slide04_explosion_cause.png']),
        ('decision/01-turn-around-or-free-return.txt', ['content/images/slide05_turn_or_free_return.png']),
        ('decision/02-freeze-or-squeeze.txt', ['content/images/slide06_cm_interior.png', 'content/images/slide06_lm_interior.png']),

        # SURVIVAL PHASE
        ('info/03-spacecraft-configuration.txt', ['content/images/slide07_spacecraft_configuration.jpg']),
        ('narrative/03-lifeboat-and-moon-flyby.txt', ['content/images/slide08_lifeboat_interior.jpg', 'content/images/slide08_moon_far_side_highres.jpg']),
        ('decision/03-stars-or-sun-earth-navigation.txt', ['content/images/slide09_aot.jpg', 'content/images/slide09_earth_terminator.jpg']),
        ('info/04-navigation-methods.txt', ['content/images/slide10_dsky.png']),
        ('decision/04-build-co2-mailbox.txt', ['content/images/slide11_mailbox.jpg']),
        ('info/05-the-mailbox.txt', []),
        ('decision/05-water-conservation.txt', ['content/images/slide13_water_rationing.png']),
        ('info/06-survival-conditions.txt', []),
        ('decision/06-speed-up-pc2-burn.txt', ['content/images/slide15_lm_descent_engine.jpg']),
        ('narrative/04-the-long-cold-journey.txt', []),
        ('decision/07-communication-discipline.txt', ['content/images/slide17_mission_control.jpg']),

        # FINAL APPROACH
        ('decision/08-battery-jump-start.txt', []),
        ('decision/09-service-module-jettison-timing.txt', ['content/images/slide19_sm_damage.jpg']),
        ('narrative/05-powering-up-and-jettison.txt', ['content/images/slide20_lm_jettison.jpg']),
        ('decision/10-computer-restart-or-manual.txt', ['content/images/slide21_computer_restart.png']),
        ('info/07-reentry-corridor-physics.txt', ['content/images/slide22_reentry_corridor.png']),
        ('narrative/06-reentry-and-splashdown.txt', ['content/images/slide23_splashdown.jpg']),
        ('info/08-lessons-learned.txt', []),
    ]

    # Process each content file
    for content_file, image_files in content_order:
        process_content_file(doc, f'content/{content_file}', image_files)
        doc.add_page_break()

    # Add final mission patch
    add_final_page(doc)

    # Save document
    output_file = 'Apollo_13_Educational_Materials.docx'
    doc.save(output_file)
    print(f"✅ Document created: {output_file}")
    return output_file

def add_title_page(doc):
    """Add a title page to the document."""
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("Apollo 13\n")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 128)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run("A Successful Failure")
    run.font.size = Pt(28)
    run.font.italic = True

    doc.add_paragraph("\n" * 3)

    # Add cover image if available
    cover_image = Path('apollo_13_scout_cover.png')
    if cover_image.exists():
        try:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
            run.add_picture(str(cover_image), width=Inches(5))
        except Exception as e:
            print(f"⚠️  Warning: Could not add cover image: {e}")

    doc.add_paragraph("\n" * 2)

    # Description
    desc = doc.add_paragraph()
    desc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = desc.add_run("Educational Materials for Boy Scouts\n")
    run.font.size = Pt(16)

    desc2 = doc.add_paragraph()
    desc2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = desc2.add_run("The story of courage, ingenuity, and teamwork\n")
    run.font.size = Pt(14)
    run.font.italic = True

    desc3 = doc.add_paragraph()
    desc3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = desc3.add_run("that brought three astronauts home from the Moon")
    run.font.size = Pt(14)
    run.font.italic = True

    doc.add_page_break()

def add_final_page(doc):
    """Add final page with mission patch."""
    heading = doc.add_heading('Apollo 13 Mission Patch', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add mission patch
    patch_image = Path('content/images/slide25_mission_patch.jpg')
    if patch_image.exists():
        try:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
            run.add_picture(str(patch_image), width=Inches(4))
        except Exception as e:
            print(f"⚠️  Warning: Could not add mission patch image: {e}")

    doc.add_paragraph()

    # Add closing text
    closing = doc.add_paragraph()
    closing.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = closing.add_run('"Ex Luna, Scientia"\n')
    run.font.size = Pt(18)
    run.font.italic = True

    closing2 = doc.add_paragraph()
    closing2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = closing2.add_run('From the Moon, Knowledge')
    run.font.size = Pt(14)
    run.font.italic = True

def process_content_file(doc, content_path, image_files):
    """Process a single content file and add it to the document."""
    content_file = Path(content_path)

    if not content_file.exists():
        print(f"⚠️  Warning: {content_path} not found")
        return

    # Read the content
    with open(content_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Parse the content
    lines = content.split('\n')

    # Extract header (first 3 lines typically)
    slide_type = lines[0].strip() if len(lines) > 0 else ""
    title = lines[1].strip() if len(lines) > 1 else ""
    subtitle = lines[2].strip() if len(lines) > 2 else ""

    # Add slide type as small header
    if slide_type:
        type_para = doc.add_paragraph(slide_type)
        type_para.runs[0].font.size = Pt(10)
        type_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # Add title
    if title:
        heading = doc.add_heading(title, level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 128)

    # Add subtitle
    if subtitle:
        sub_para = doc.add_paragraph(subtitle)
        sub_para.runs[0].font.size = Pt(12)
        sub_para.runs[0].font.italic = True
        sub_para.runs[0].font.color.rgb = RGBColor(60, 60, 60)

    doc.add_paragraph()

    # Add images at the top if available
    for image_file in image_files:
        image_path = Path(image_file)
        if image_path.exists():
            try:
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run()
                # Adjust size based on image - max width 6 inches
                run.add_picture(str(image_path), width=Inches(6))
                doc.add_paragraph()
            except Exception as e:
                print(f"⚠️  Warning: Could not add image {image_file}: {e}")
                # Add a note that the image couldn't be loaded
                p = doc.add_paragraph(f"[Image: {image_path.name}]")
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.runs[0].font.italic = True

    # Process sections
    current_section = None
    in_section = False

    for line in lines[3:]:  # Skip the first 3 header lines
        line = line.rstrip()

        # Check for section markers
        if line.startswith('===') and line.endswith('==='):
            section_name = line.strip('= ').strip()
            current_section = doc.add_heading(section_name, level=2)
            current_section.runs[0].font.color.rgb = RGBColor(0, 64, 128)
            in_section = True
            continue

        # Skip empty lines at the start
        if not line and not in_section:
            continue

        # Process content lines
        if line:
            # Check if it's a bullet point
            if line.startswith('•') or line.startswith('▸'):
                p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
            elif line.startswith('✅') or line.startswith('❌') or line.startswith('⏳') or line.startswith('❓'):
                p = doc.add_paragraph(line.strip())
            # Check for emphasis (lines that are short and might be headers)
            elif line.isupper() and len(line) < 50:
                p = doc.add_paragraph(line)
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(11)
            # Check for special formatting (PROS:, CONS:, etc.)
            elif line.endswith(':') and len(line) < 40:
                p = doc.add_paragraph(line)
                p.runs[0].font.bold = True
            else:
                # Regular paragraph
                doc.add_paragraph(line)
        else:
            # Add spacing between paragraphs
            doc.add_paragraph()

if __name__ == '__main__':
    create_apollo13_document()
